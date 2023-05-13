import docx
import os

doc = docx.Document()

doc.add_paragraph('Hola, esto es un documento de Word creado desde Python.')
doc.add_heading('Esto es un título', level=1)
doc.add_paragraph('Este es un párrafo con más texto.')

# Guardar el documento
doc.save('mi_documento.docx')

doc = docx.Document('mi_documento.docx')

# Recorrer los párrafos del documento y obtener su contenido
for para in doc.paragraphs:
    print(para.text)

# Recorrer las tablas del documento y obtener su contenido
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            print(cell.text)
            
doc = docx.Document('mi_documento.docx')

# Agregar una imagen al documento
doc.add_picture('ruta/a/mi/imagen.jpg', width=docx.shared.Cm(10), height=docx.shared.Cm(10))

# Guardar el documento actualizado
doc.save('mi_documento_actualizado.docx')


doc = docx.Document('mi_documento.docx')

# Agregar un nuevo párrafo al final del documento
doc.add_paragraph('Este es un nuevo párrafo agregado al final del documento.')

# Guardar el documento actualizado en el mismo archivo
doc.save('mi_documento.docx')

doc = docx.Document('mi_documento.docx')

# Seleccionar el segundo párrafo
para = doc.paragraphs[1]

# Actualizar el contenido del párrafo
para.text = 'Este es un nuevo contenido para el segundo párrafo.'

# Guardar el documento actualizado en el mismo archivo
doc.save('mi_documento.docx')


os.remove("archivo.docx")